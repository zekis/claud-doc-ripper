#!/usr/bin/env python3
"""
Knowledge Base Builder - Builds cumulative knowledge base from Word documents
Now using Kimi 2.5 via Moonshot API (OpenAI-compatible)

Organizes extracted information into:
- Products/[ProductName]/ - Product knowledge that accumulates over time
- Clients/[ClientName]/ - Client information split into categories
- Reference Materials - Guides, templates, procedures, and other reusable content organized by product and type
"""

import json
import argparse
import os
from pathlib import Path
from typing import List, Callable, Any, get_origin, get_args, Union
from docx import Document
from openai import OpenAI
import inspect
import datetime

# Load environment variables from .env file if it exists
_ENV_LOADED = False
try:
    from dotenv import load_dotenv
    env_path = Path(__file__).parent / '.env'
    _ENV_LOADED = load_dotenv(dotenv_path=env_path, verbose=True)
    if _ENV_LOADED:
        print(f"[DEBUG] .env file loaded from: {env_path}")
    else:
        print(f"[DEBUG] .env file not found at: {env_path}")
except ImportError:
    # python-dotenv not installed, will use system environment variables
    print("[DEBUG] python-dotenv not installed, using system environment variables")
    pass

# Configuration
MOONSHOT_API_KEY = os.getenv("MOONSHOT_API_KEY")
MOONSHOT_BASE_URL = os.getenv("MOONSHOT_BASE_URL", "https://api.moonshot.ai/v1")  # Official Moonshot API endpoint

# Company Context (for better AI understanding)
COMPANY_NAME = os.getenv("COMPANY_NAME", "")
COMPANY_FORMER_NAME = os.getenv("COMPANY_FORMER_NAME", "")
COMPANY_BUSINESS = os.getenv("COMPANY_BUSINESS", "")
COMPANY_INDUSTRIES = os.getenv("COMPANY_INDUSTRIES", "")

print(f"[DEBUG] After loading - API key present: {bool(MOONSHOT_API_KEY)}")

# Model selection (Kimi models - based on official API documentation)
MODEL_FAST = "kimi-k2.5"      # Kimi K2.5 model for fast extraction
MODEL_SMART = "kimi-k2.5"     # Kimi K2.5 model for complex analysis

# Global variable to store current document data for retrieval functions
_current_document_data = None


def _is_optional(annotation) -> bool:
    """
    Check if a type annotation is Optional (Union with None).

    Args:
        annotation: The type annotation to check

    Returns:
        True if the annotation is Optional[T] (i.e., Union[T, None])
    """
    if annotation == inspect.Parameter.empty:
        return False

    # Get the origin type (e.g., Union from Union[str, None])
    origin = get_origin(annotation)

    # Check if it's a Union type
    if origin is Union:
        # Get the types in the Union
        args = get_args(annotation)
        # Check if None is one of the types
        return type(None) in args

    return False


def _get_type_schema(annotation) -> dict:
    """
    Convert a Python type annotation to OpenAI function parameter schema.

    Args:
        annotation: The type annotation from inspect.signature

    Returns:
        Dict with 'type' and optionally 'items' for arrays
    """
    if annotation == inspect.Parameter.empty:
        return {"type": "string"}

    # Handle basic types
    if annotation == int:
        return {"type": "integer"}
    elif annotation == float:
        return {"type": "number"}
    elif annotation == bool:
        return {"type": "boolean"}
    elif annotation == str:
        return {"type": "string"}
    elif annotation == dict:
        return {"type": "object"}
    elif annotation == list:
        return {"type": "array"}

    # Handle typing module types
    origin = get_origin(annotation)
    args = get_args(annotation)

    # Handle List[T]
    if origin is list:
        if args:
            # Get the inner type
            inner_type = _get_type_schema(args[0])
            return {
                "type": "array",
                "items": inner_type
            }
        return {"type": "array"}

    # Handle Union types (including Optional)
    if origin is Union:
        # Filter out None from the union
        non_none_types = [arg for arg in args if arg != type(None)]

        if len(non_none_types) == 1:
            # This is Optional[T], return schema for T
            return _get_type_schema(non_none_types[0])
        else:
            # Multiple non-None types - can't represent in JSON schema easily
            # Default to the first type
            return _get_type_schema(non_none_types[0])

    # Default to string for unknown types
    return {"type": "string"}


def parse_json_response(content: str) -> Any:
    """
    Parse JSON from AI response, handling markdown code blocks.

    Args:
        content: The raw response content that may contain JSON wrapped in markdown

    Returns:
        Parsed JSON as dict, list, or other JSON-compatible type

    Raises:
        json.JSONDecodeError: If the content is not valid JSON
    """
    if not content:
        raise json.JSONDecodeError("Empty content", "", 0)

    # Strip whitespace
    cleaned = content.strip()

    # Look for markdown code fence with json language identifier
    if '```json' in cleaned:
        # Extract content between ```json and closing ```
        start = cleaned.find('```json') + 7  # len('```json')
        end = cleaned.find('```', start)
        if end != -1:
            cleaned = cleaned[start:end].strip()
        else:
            # No closing fence, take everything after ```json
            cleaned = cleaned[start:].strip()
    # Look for generic markdown code fence at start
    elif cleaned.startswith('```'):
        # Remove opening fence
        cleaned = cleaned.split('\n', 1)[1] if '\n' in cleaned else cleaned[3:]
        if cleaned.endswith('```'):
            # Remove closing fence
            cleaned = cleaned.rsplit('\n', 1)[0] if '\n' in cleaned else cleaned[:-3]
        cleaned = cleaned.strip()

    # If still has explanatory text before JSON, find JSON start
    if cleaned and cleaned[0] not in ('{', '['):
        # Find first { or [ which indicates JSON start
        json_start = -1
        for char in ('{', '['):
            pos = cleaned.find(char)
            if pos != -1 and (json_start == -1 or pos < json_start):
                json_start = pos

        if json_start != -1:
            cleaned = cleaned[json_start:]
        else:
            # No JSON found
            raise json.JSONDecodeError(f"No JSON found in response. Content starts with: {cleaned[:100]}", "", 0)

    return json.loads(cleaned)


def function_to_tool_schema(func: Callable) -> dict:
    """Convert a Python function with type hints and docstring to OpenAI tool schema"""
    sig = inspect.signature(func)
    doc = inspect.getdoc(func) or ""
    
    # Parse docstring for parameter descriptions
    param_descriptions = {}
    lines = doc.split('\n')
    in_args_section = False
    for line in lines:
        line = line.strip()
        if line.startswith('Args:'):
            in_args_section = True
            continue
        if in_args_section and ':' in line:
            param_name = line.split(':')[0].strip()
            param_desc = ':'.join(line.split(':')[1:]).strip()
            param_descriptions[param_name] = param_desc
        elif in_args_section and line and not line[0].isspace():
            in_args_section = False
    
    # Build parameters schema
    properties = {}
    required = []

    for param_name, param in sig.parameters.items():
        if param_name == 'self':
            continue

        param_schema = _get_type_schema(param.annotation)
        param_schema["description"] = param_descriptions.get(param_name, f"The {param_name} parameter")

        properties[param_name] = param_schema

        # Check if parameter is required (no default value and not Optional)
        if param.default == inspect.Parameter.empty and not _is_optional(param.annotation):
            required.append(param_name)
    
    # Get function description from docstring
    description_lines = []
    for line in lines:
        if line.startswith('Args:') or line.startswith('Returns:'):
            break
        if line:
            description_lines.append(line)
    description = ' '.join(description_lines) or func.__name__
    
    return {
        "type": "function",
        "function": {
            "name": func.__name__,
            "description": description,
            "parameters": {
                "type": "object",
                "properties": properties,
                "required": required
            }
        }
    }


def run_with_tools(client: OpenAI, model: str, prompt: str, functions: List[Callable] = None, 
                   return_type: type = str, timeout: int = 300) -> Any:
    """
    Run a chat completion with optional function calling support.
    
    This mimics the Auggie SDK's agent.run() interface but uses OpenAI/Moonshot API.
    """
    messages = [{"role": "user", "content": prompt}]
    
    # Convert functions to tool schemas
    tools = None
    tool_map = {}
    if functions:
        tools = [function_to_tool_schema(func) for func in functions]
        tool_map = {func.__name__: func for func in functions}
    
    max_iterations = 20  # Prevent infinite loops
    iteration = 0
    
    while iteration < max_iterations:
        iteration += 1
        
        # Make API call
        completion_args = {
            "model": model,
            "messages": messages,
            "temperature": 1,  # Kimi K2.5 requires temperature=1
        }
        
        if tools:
            completion_args["tools"] = tools
            completion_args["tool_choice"] = "auto"
        
        response = client.chat.completions.create(**completion_args)
        choice = response.choices[0]
        
        # Check if we're done
        if choice.finish_reason == "stop":
            content = choice.message.content
            # Try to parse as JSON if return_type is not str
            if return_type != str and content:
                try:
                    return parse_json_response(content)
                except:
                    return content
            return content
        
        # Handle tool calls
        if choice.finish_reason == "tool_calls" and choice.message.tool_calls:
            messages.append(choice.message)
            
            for tool_call in choice.message.tool_calls:
                function_name = tool_call.function.name
                function_args = json.loads(tool_call.function.arguments)

                # Execute the function with error handling
                if function_name in tool_map:
                    try:
                        function_result = tool_map[function_name](**function_args)
                    except Exception as e:
                        # If function execution fails, return error to AI
                        import traceback
                        error_msg = f"Error executing {function_name}: {str(e)}\n{traceback.format_exc()[:200]}"
                        function_result = {"error": error_msg}
                        print(f"      [ERROR] Tool execution failed: {error_msg}")

                    # Add function result to messages
                    messages.append({
                        "role": "tool",
                        "tool_call_id": tool_call.id,
                        "name": function_name,
                        "content": json.dumps(function_result) if isinstance(function_result, dict) else str(function_result)
                    })
                else:
                    # Function not found in tool map
                    messages.append({
                        "role": "tool",
                        "tool_call_id": tool_call.id,
                        "name": function_name,
                        "content": json.dumps({"error": f"Function {function_name} not found"})
                    })
        else:
            # Unexpected finish reason
            break
    
    # If we get here, return the last message content
    return choice.message.content if choice.message.content else ""


# ============================================================================
# Document Structure Extraction
# ============================================================================

def extract_document_structure(doc_path: str) -> dict:
    """Extract document structure (headings hierarchy) from Word document

    Returns:
        dict with:
            - 'structure': List of headings with their levels and indices
            - 'sections': Dict mapping section indices to their content
            - 'full_content': Full document content
            - 'metadata': Document metadata (author, created, modified, etc.)
    """
    print(f"[Extracting] Document structure from: {doc_path}")
    doc = Document(doc_path)

    # Extract document metadata
    doc_props = doc.core_properties
    doc_metadata = {
        "author": doc_props.author or "Unknown",
        "created": doc_props.created.isoformat() if doc_props.created else None,
        "modified": doc_props.modified.isoformat() if doc_props.modified else None,
        "last_modified_by": doc_props.last_modified_by or "Unknown",
        "revision": doc_props.revision,
        "title": doc_props.title or Path(doc_path).stem,
        "subject": doc_props.subject or "",
        "keywords": doc_props.keywords or "",
    }

    structure = []
    sections = {}
    current_section = {"index": 0, "heading": "Introduction", "level": 0, "content": []}
    section_index = 0

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        style = para.style.name
        text = para.text.strip()

        # Check if this is a heading
        if 'Heading' in style:
            # Save previous section
            if current_section["content"]:
                sections[current_section["index"]] = {
                    "heading": current_section["heading"],
                    "level": current_section["level"],
                    "content": "\n".join(current_section["content"])
                }

            # Start new section
            section_index += 1
            level = 1
            if 'Heading 1' in style:
                level = 1
            elif 'Heading 2' in style:
                level = 2
            elif 'Heading 3' in style:
                level = 3
            elif 'Heading 4' in style:
                level = 4
            else:
                level = 5

            current_section = {
                "index": section_index,
                "heading": text,
                "level": level,
                "content": []
            }

            structure.append({
                "index": section_index,
                "heading": text,
                "level": level
            })
        else:
            # Add to current section
            current_section["content"].append(text)

    # Save last section
    if current_section["content"]:
        sections[current_section["index"]] = {
            "heading": current_section["heading"],
            "level": current_section["level"],
            "content": "\n".join(current_section["content"])
        }

    # Extract tables and add to sections
    for table in doc.tables:
        table_content = ["[TABLE]"]
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            table_content.append(row_text)
        table_content.append("[/TABLE]")

        # Add table to the last section (or create a Tables section if needed)
        if section_index in sections:
            sections[section_index]["content"] += "\n\n" + "\n".join(table_content)

    # Build full content
    full_content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    print(f"[OK] Found {len(structure)} sections")
    print(f"[Metadata] Author: {doc_metadata['author']}, Modified: {doc_metadata['modified'] or 'Unknown'}\n")

    return {
        "structure": structure,
        "sections": sections,
        "full_content": full_content,
        "metadata": doc_metadata
    }


# ============================================================================
# Document Retrieval Functions (called by AI as tools)
# ============================================================================

def get_section_by_index(section_index: int) -> str:
    """Retrieve a specific section from the document by its index.

    Args:
        section_index: The index number of the section to retrieve (starting from 1)

    Returns:
        The content of the requested section including its heading
    """
    global _current_document_data

    import datetime
    timestamp = datetime.datetime.now().strftime("%H:%M:%S")
    print(f"      [{timestamp}] [TOOL CALL] get_section_by_index({section_index})")

    if _current_document_data is None:
        return "Error: No document is currently loaded"

    sections = _current_document_data.get("sections", {})

    if section_index not in sections:
        available = ", ".join(str(i) for i in sorted(sections.keys()))
        return f"Error: Section {section_index} not found. Available sections: {available}"

    section = sections[section_index]
    print(f"      [{timestamp}] [RETRIEVED] Section {section_index}: {section['heading']}")
    return f"# {section['heading']}\n\n{section['content']}"


def get_section_by_heading(heading_keyword: str) -> str:
    """Search for and retrieve a section by heading name or keyword.

    Args:
        heading_keyword: A keyword or phrase to search for in section headings

    Returns:
        The content of the first matching section, or list of matches if multiple found
    """
    global _current_document_data

    import datetime
    timestamp = datetime.datetime.now().strftime("%H:%M:%S")
    print(f"      [{timestamp}] [TOOL CALL] get_section_by_heading('{heading_keyword}')")

    if _current_document_data is None:
        return "Error: No document is currently loaded"

    sections = _current_document_data.get("sections", {})
    structure = _current_document_data.get("structure", [])

    # Find matching sections
    matches = []
    for item in structure:
        if heading_keyword.lower() in item["heading"].lower():
            matches.append(item)

    if not matches:
        return f"Error: No sections found matching '{heading_keyword}'"

    if len(matches) == 1:
        idx = matches[0]["index"]
        section = sections[idx]
        print(f"      [{timestamp}] [RETRIEVED] Section {idx}: {section['heading']}")
        return f"# {section['heading']}\n\n{section['content']}"
    else:
        # Multiple matches - return list
        match_list = "\n".join([f"  {m['index']}. {m['heading']}" for m in matches])
        print(f"      [{timestamp}] [RETRIEVED] {len(matches)} matching sections")
        return f"Multiple sections found matching '{heading_keyword}':\n{match_list}\n\nUse get_section_by_index() to retrieve a specific one."


def get_multiple_sections(section_indices: list) -> str:
    """Retrieve multiple sections at once by their indices.

    Args:
        section_indices: List of section index numbers to retrieve

    Returns:
        Combined content of all requested sections
    """
    global _current_document_data

    import datetime
    timestamp = datetime.datetime.now().strftime("%H:%M:%S")
    print(f"      [{timestamp}] [TOOL CALL] get_multiple_sections({section_indices})")

    if _current_document_data is None:
        return "Error: No document is currently loaded"

    sections = _current_document_data.get("sections", {})
    result = []

    retrieved_count = 0
    for idx in section_indices:
        if idx in sections:
            section = sections[idx]
            result.append(f"# {section['heading']}\n\n{section['content']}\n")
            retrieved_count += 1
        else:
            result.append(f"[Section {idx} not found]\n")

    print(f"      [{timestamp}] [RETRIEVED] {retrieved_count}/{len(section_indices)} sections")
    return "\n---\n\n".join(result)


# ============================================================================
# Helper Functions
# ============================================================================

def get_company_context() -> str:
    """Build company context string for AI prompts"""
    context_parts = []

    if COMPANY_NAME:
        context_parts.append(f"COMPANY: {COMPANY_NAME}")
        if COMPANY_FORMER_NAME:
            context_parts.append(f"  (formerly {COMPANY_FORMER_NAME})")

    if COMPANY_BUSINESS:
        context_parts.append(f"BUSINESS: {COMPANY_BUSINESS}")

    if COMPANY_INDUSTRIES:
        context_parts.append(f"INDUSTRIES: {COMPANY_INDUSTRIES}")

    if context_parts:
        return "COMPANY CONTEXT:\n" + "\n".join(context_parts) + "\n\n"
    return ""


# ============================================================================
# Extraction Functions
# ============================================================================

def extract_metadata(content: str, client: OpenAI, existing_structure: str) -> dict:
    """Extract metadata: products mentioned, client name, document type"""
    print("[Extracting] Document metadata...")

    sample = content[:15000]

    company_context = get_company_context()

    try:
        result = run_with_tools(
            client=client,
            model=MODEL_FAST,
            prompt=f"""{company_context}Analyze this document intelligently and extract all relevant metadata.

            {existing_structure}

            DOCUMENT CONTENT:
            {sample}

            INSTRUCTIONS:
            1. **Products/Systems**: Identify products/systems that are SUBSTANTIALLY discussed
               - ONLY include if: The document provides technical details, procedures, or configuration info
               - EXCLUDE if: Just mentioned in passing, listed in a table, or used as an example
               - Look for: Product names, software systems, equipment models, platforms, tools
               - Examples: "BULKmetrix", "QMS", "Insight CM", "Azure DevOps", "Tortoise Git"
               - If products already exist in structure above, use EXACT same name
               - Quality over quantity - better to miss a minor mention than create useless stubs

            2. **Client**: Identify the client/customer this document is for
               - Look for: Company names, project names, client references
               - Distinguish between: our company ({COMPANY_NAME or 'us'}), the client, and vendors
               - If client exists in structure above, use EXACT same name
               - Return null if this is internal documentation (no specific client)

            3. **Document Type**: What kind of document is this?
               - Examples: "User Manual", "Technical Specification", "How-To Guide",
                 "Installation Guide", "Configuration Guide", "Process Document",
                 "Technical Guide", "Reference Manual", "Standard Operating Procedure"

            4. **Document Category**: What technical area does this cover?
               - Examples: "Version Control", "Controls Systems", "Electrical",
                 "Installation", "Configuration", "Maintenance", "Safety", "Engineering"

            Return ONLY valid JSON in this exact format:
            {{
                "products": ["Product1", "Product2"],
                "client_name": "ClientName or null",
                "document_type": "DocumentType",
                "document_category": "Category"
            }}
            """,
            return_type=str,
            timeout=120
        )

        # Parse JSON
        metadata = parse_json_response(result)
        print(f"[OK] Found: {len(metadata.get('products', []))} products, Client: {metadata.get('client_name', 'Unknown')}\n")
        return metadata

    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"[ERROR] Failed to extract metadata: {e}")
        print(f"   Details: {error_details[:500]}...\n")
        return {
            "products": [],
            "client_name": None,
            "document_type": "Unknown",
            "document_category": "General"
        }


def extract_product_knowledge(product_name: str, doc_data: dict, client: OpenAI, existing_structure: str) -> str:
    """Extract knowledge about a specific product using tool-based retrieval"""
    global _current_document_data
    _current_document_data = doc_data

    print(f"   [Product] Extracting knowledge for: {product_name}")

    # Build structure summary for the AI
    structure = doc_data.get("structure", [])
    structure_summary = "DOCUMENT STRUCTURE:\n"
    for item in structure:
        indent = "  " * (item["level"] - 1)
        structure_summary += f"{indent}{item['index']}. {item['heading']} (Level {item['level']})\n"

    company_context = get_company_context()

    try:
        result = run_with_tools(
            client=client,
            model=MODEL_SMART,
            prompt=f"""{company_context}You are analyzing a technical document to extract knowledge about {product_name}.

            {existing_structure}

            NOTE: If {product_name} already exists in the structure above, this content will be ADDED to existing knowledge.
            Focus on extracting NEW information that complements what might already exist.

            {structure_summary}

            TASK: Write a professional wiki article about {product_name}.

            CRITICAL - INSUFFICIENT INFORMATION DETECTION:
            - If {product_name} is only mentioned in passing (1-2 brief mentions)
            - If there's no technical information about {product_name}
            - If {product_name} is just listed as an example or in a table
            - Then return ONLY this text: "INSUFFICIENT_INFORMATION"
            - Do NOT create a stub article - just return the marker

            CRITICAL - WIKI STYLE REQUIREMENTS (when sufficient information exists):
            - Write DIRECTLY as a wiki article - ABSOLUTELY NO meta-commentary
            - NEVER say: "after reviewing", "based on the document", "this document", "let me", "I now have", "I will write", etc.
            - Do NOT explain your analysis process or what you're doing
            - Start IMMEDIATELY with "## Overview" - no preamble text before it
            - Write in present tense, factual, encyclopedic style
            - Imagine you are Wikipedia - direct facts only

            WRONG (DO NOT DO THIS):
            "Based on the document, I now have sufficient information to write about Product X."

            RIGHT (DO THIS):
            "## Overview\n\nProduct X is a software platform that..."

            Create a wiki article with these sections (only include sections with actual content):

            ## Overview
            What {product_name} is, its purpose, and how it's used in our industry.

            ## Features & Capabilities
            List features, technical specifications, and key functionalities.

            ## Integration & Interfaces
            How it integrates with other systems, data exchanges, connected systems.

            ## Configuration & Setup
            Configuration options, settings, setup requirements, installation.

            ## Usage & Operations
            How to use it, common operations, workflows, procedures.

            ## Technical Details
            Architecture, design, database schema, APIs, technical specs.

            ## Engineering Notes
            How we configure/customize it, common issues, solutions, best practices.

            INSTRUCTIONS:
            1. Use the tools to retrieve sections about {product_name}
            2. Write clean, direct wiki content
            3. Use present tense (e.g., "Git is...", not "The document describes Git as...")
            4. Include tables, lists, code blocks where appropriate
            5. Be specific and technical - avoid generic statements
            6. If minimal information exists, write brief article and stop

            CRITICAL: Your response must START with "## Overview" immediately.
            NO text before it. NO explanations. NO "Based on...". NO "I will...".
            Return ONLY pure wiki article markdown - start with ## Overview heading.
            """,
            functions=[get_section_by_index, get_section_by_heading, get_multiple_sections],
            return_type=str,
            timeout=300
        )

        print(f"   [OK] Extracted {len(result)} characters of knowledge\n")
        return result

    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"   [ERROR] Failed to extract product knowledge: {e}")
        print(f"   Details: {error_details[:500]}...\n")
        return f"# {product_name}\n\nError extracting knowledge: {e}\n\nDetails:\n{error_details}"

    finally:
        # Always clean up global state
        _current_document_data = None


def extract_document_template(product_name: str, doc_data: dict, client: OpenAI, existing_structure: str) -> list:
    """Extract reusable reference materials from document using tool-based retrieval"""
    global _current_document_data
    _current_document_data = doc_data

    print(f"   [Analyzing] Document for reusable reference materials: {product_name}")

    # Build structure summary for the AI
    structure = doc_data.get("structure", [])
    structure_summary = "DOCUMENT STRUCTURE:\n"
    for item in structure:
        indent = "  " * (item["level"] - 1)
        structure_summary += f"{indent}{item['index']}. {item['heading']} (Level {item['level']})\n"

    result = None
    company_context = get_company_context()

    try:
        result = run_with_tools(
            client=client,
            model=MODEL_SMART,
            prompt=f"""{company_context}We are building a comprehensive technical knowledge base. Analyze this document intelligently and extract ALL valuable knowledge for future reference.

            {existing_structure}

            {structure_summary}

            TASK: Identify 1-3 different ways this document provides valuable knowledge. Think beyond just the document itself - extract the KNOWLEDGE it contains.

            INSTRUCTIONS:
            1. Review the document structure above
            2. Use the provided tools to retrieve relevant sections
            3. Extract knowledge that would be useful to engineers, technicians, and project teams

            KNOWLEDGE CATEGORIES TO CONSIDER:

            **DOCUMENT_TEMPLATE** - Document structure/format for creating similar documents
            - Extract: Document structure, required sections, formatting patterns
            - Use when: The document itself is a good template for future docs

            **ENGINEERING** - Technical configuration, setup, or design knowledge
            - Extract: How systems are configured, designed, or integrated
            - Examples: "Product X is configured by...", "System Y integrates with..."
            - Use for: Engineering knowledge about products/systems

            **HOW_TO** - Step-by-step instructions to accomplish specific tasks
            - Extract: Procedural knowledge, workflows, operational steps
            - Examples: "How to deploy code", "How to configure Git"
            - Use for: Instructional guides

            **INSTALLATION** - Installation and setup procedures
            - Extract: Installation steps, setup requirements, pre-requisites
            - Use for: Installing systems, software, equipment

            **CONFIGURATION** - Configuration guides and settings
            - Extract: Configuration options, settings, parameters
            - Use for: Setting up and configuring systems

            **SETUP** - Initial setup and initialization procedures
            - Extract: Setup workflows, initialization steps
            - Use for: Getting systems ready for use

            **PROCEDURE** - Standard operating procedures or processes
            - Extract: Formal processes, workflows, standard methods
            - Use for: Standardized operational procedures

            **REFERENCE** - Reference information, standards, specifications
            - Extract: Technical specs, standards, lookup information
            - Use for: Quick reference material

            **BEST_PRACTICES** - Recommended approaches and methodologies
            - Extract: Best practices, recommendations, lessons learned
            - Use for: Guidance on proper methods

            For EACH piece of valuable knowledge, return a JSON object with:
            - type: One of the types above (DOCUMENT_TEMPLATE, ENGINEERING, HOW_TO, etc.)
            - title: Clear, descriptive title (e.g., "How to Configure Git Repositories", "BULKmetrix Configuration Guide")
            - category: Technical category (e.g., "Version Control", "Controls Systems", "Installation")
            - content: The extracted knowledge in complete markdown format (not just a summary!)
            - tags: 3-5 relevant tags (lowercase, hyphenated: e.g., "git", "version-control", "azure-devops")

            Return a JSON array of 1-3 knowledge items:
            [
                {{
                    "type": "HOW_TO",
                    "title": "How to Configure Git Repositories",
                    "category": "Version Control",
                    "content": "# How to Configure Git Repositories\\n\\n## Overview\\n...full content...",
                    "tags": ["git", "version-control", "configuration", "repository"]
                }}
            ]

            CRITICAL - WIKI CONTENT REQUIREMENTS:
            - Write content as clean, professional wiki articles
            - ABSOLUTELY NO meta-commentary in the content field
            - NEVER say: "this guide covers", "we will show", "based on the document", "I now have", etc.
            - Start directly with the heading (e.g., "# How to Install Git\\n\\nGit is installed by...")
            - NO explanatory text before the heading - start immediately with #
            - Use present tense, factual, encyclopedic style
            - Extract COMPLETE content, not summaries
            - Be specific to {product_name} where relevant
            - Include code blocks, tables, lists as appropriate

            WRONG (DO NOT put this in content):
            "Based on the document, here is how to install Git..."

            RIGHT (DO put this in content):
            "# How to Install Git\\n\\nGit is installed by downloading..."

            Return ONLY valid JSON array, no preamble or explanation before it.
            """,
            functions=[get_section_by_index, get_section_by_heading, get_multiple_sections],
            return_type=list,
            timeout=300
        )

        # Parse JSON if it's a string
        if isinstance(result, str):
            result = parse_json_response(result)

        if not isinstance(result, list):
            result = [result]

        print(f"   [OK] Identified {len(result)} reference material(s)\n")
        return result

    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"   [ERROR] Failed to extract reference materials: {e}")
        print(f"   Details: {error_details[:500]}...")
        if result:
            print(f"   [DEBUG] Raw result: {result[:500]}...\n")
        return []

    finally:
        # Always clean up global state
        _current_document_data = None


def extract_client_info(client_name: str, content: str, client: OpenAI, existing_structure: str) -> dict:
    """Extract client information split into categories"""
    # Skip if no valid client name
    if not client_name or client_name.lower() in ['none', 'unknown', 'n/a']:
        return {}

    print(f"   [Client] Extracting information for: {client_name}")

    sample = content[:20000]

    try:
        result = run_with_tools(
            client=client,
            model=MODEL_FAST,
            prompt=f"""Extract information about {client_name} from this document and organize into categories.

            {existing_structure}

            NOTE: Check if {client_name} already exists in the structure above.
            If they do, check what categories already exist (e.g., overview.md, locations.md, hardware.md).
            You can use existing categories OR suggest new ones that fit the pattern.

            DOCUMENT CONTENT:
            {sample}

            Return ONLY valid JSON in this exact format:
            {{
                "overview": "Brief overview of the client and project",
                "locations": ["Location 1 with details", "Location 2 with details"],
                "hardware": ["Hardware item 1", "Hardware item 2"],
                "configuration": ["Config detail 1", "Config detail 2"],
                "contacts": ["Contact 1", "Contact 2"]
            }}

            You can add additional categories if needed (e.g., "software", "network", "security").
            Be comprehensive and extract all relevant details.
            """,
            return_type=str,
            timeout=120
        )

        # Parse JSON
        client_data = parse_json_response(result)
        print(f"   [OK] Extracted {len(client_data)} categories\n")
        return client_data

    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"   [ERROR] Failed to extract client info: {e}")
        print(f"   Details: {error_details[:500]}...\n")
        return {}


# ============================================================================
# File Saving Functions
# ============================================================================

def scan_existing_structure(base_dir: str) -> str:
    """Scan existing knowledge base structure to help AI maintain consistency"""
    base_path = Path(base_dir)
    if not base_path.exists():
        return "EXISTING KNOWLEDGE BASE STRUCTURE: Empty (this is the first document)"

    structure_info = ["EXISTING KNOWLEDGE BASE STRUCTURE:"]

    # Scan Products
    products_dir = base_path / "Products"
    if products_dir.exists():
        structure_info.append("\nProducts:")
        for product_dir in sorted(products_dir.iterdir()):
            if product_dir.is_dir():
                structure_info.append(f"  - {product_dir.name}")
                # Check for reference materials
                ref_dir = product_dir / "Reference Materials"
                if ref_dir.exists():
                    for category_dir in sorted(ref_dir.iterdir()):
                        if category_dir.is_dir():
                            structure_info.append(f"      Reference Materials/{category_dir.name}/")

    # Scan Clients
    clients_dir = base_path / "Clients"
    if clients_dir.exists():
        structure_info.append("\nClients:")
        for client_dir in sorted(clients_dir.iterdir()):
            if client_dir.is_dir():
                structure_info.append(f"  - {client_dir.name}")
                # List existing files
                files = [f.name for f in client_dir.iterdir() if f.is_file()]
                if files:
                    structure_info.append(f"      Files: {', '.join(files)}")

    return "\n".join(structure_info)


def save_product_knowledge(base_dir: Path, product_name: str, knowledge: str,
                           reference_materials: list, doc_type: str, doc_category: str, doc_metadata: dict = None):
    """Save product knowledge and reference materials"""
    from datetime import datetime
    import re

    product_dir = base_dir / "Products" / product_name
    product_dir.mkdir(parents=True, exist_ok=True)

    # Save product knowledge with YAML front matter
    knowledge_file = product_dir / "overview.md"

    # Check if file exists and if we should update it
    should_update = True
    if knowledge_file.exists() and doc_metadata and doc_metadata.get('modified'):
        # Read existing file to check last document date
        try:
            with open(knowledge_file, 'r', encoding='utf-8') as f:
                content = f.read()
                # Extract source_document_modified from YAML
                match = re.search(r'source_document_modified:\s*"?([^"\n]+)"?', content)
                if match:
                    existing_date = match.group(1)
                    new_date = doc_metadata.get('modified')
                    if existing_date >= new_date:
                        print(f"      [SKIP] {knowledge_file.name} - existing article is up-to-date (source: {existing_date})")
                        should_update = False
        except Exception:
            # If error reading existing file, proceed with update
            pass

    if should_update:
        with open(knowledge_file, 'w', encoding='utf-8') as f:
            # YAML front matter
            f.write("---\n")
            f.write(f"title: \"{product_name}\"\n")
            f.write(f"type: \"Product Overview\"\n")
            f.write(f"product: \"{product_name}\"\n")
            f.write(f"date_updated: \"{datetime.now().strftime('%Y-%m-%d')}\"\n")

            # Add source document metadata
            if doc_metadata:
                if doc_metadata.get('author'):
                    f.write(f"source_document_author: \"{doc_metadata['author']}\"\n")
                if doc_metadata.get('modified'):
                    f.write(f"source_document_modified: \"{doc_metadata['modified']}\"\n")
                if doc_metadata.get('title'):
                    f.write(f"source_document_title: \"{doc_metadata['title']}\"\n")

            f.write("---\n\n")

            f.write(f"# {product_name}\n\n")
            if doc_metadata and doc_metadata.get('modified'):
                f.write(f"*Source document last modified: {doc_metadata['modified'][:10]}*\n\n")
            f.write(knowledge)
        print(f"      [OK] Saved: {knowledge_file}")

    # Save each reference material (guides, templates, procedures, etc.)
    for ref_material in reference_materials:
        ref_type = ref_material.get('type', 'REFERENCE')
        ref_title = ref_material.get('title', doc_type)
        ref_category = ref_material.get('category', doc_category)
        ref_content = ref_material.get('content', '')
        ref_tags = ref_material.get('tags', [])

        # Create safe filename from title
        safe_filename = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in ref_title)
        safe_filename = safe_filename.replace(' ', '_')

        reference_dir = product_dir / "Reference Materials" / ref_category
        reference_dir.mkdir(parents=True, exist_ok=True)
        reference_file = reference_dir / f"{safe_filename}.md"

        with open(reference_file, 'w', encoding='utf-8') as f:
            # YAML front matter
            f.write("---\n")
            f.write(f"title: \"{ref_title}\"\n")
            f.write(f"type: \"{ref_type}\"\n")
            f.write(f"category: \"{ref_category}\"\n")
            f.write(f"product: \"{product_name}\"\n")
            f.write(f"source_document: \"{doc_type}\"\n")
            f.write(f"date_extracted: \"{datetime.now().strftime('%Y-%m-%d')}\"\n")

            # Add source document metadata
            if doc_metadata:
                if doc_metadata.get('author'):
                    f.write(f"source_document_author: \"{doc_metadata['author']}\"\n")
                if doc_metadata.get('modified'):
                    f.write(f"source_document_modified: \"{doc_metadata['modified']}\"\n")

            # Tags as YAML array
            if ref_tags:
                f.write("tags:\n")
                for tag in ref_tags:
                    f.write(f"  - {tag}\n")
            else:
                f.write("tags: []\n")

            f.write("---\n\n")

            # Content
            f.write(ref_content)
        print(f"      [OK] Saved: {reference_file}")


def save_client_info(base_dir: Path, client_name: str, client_data: dict):
    """Save client information split into separate files"""
    from datetime import datetime

    client_dir = base_dir / "Clients" / client_name
    client_dir.mkdir(parents=True, exist_ok=True)

    date_str = datetime.now().strftime('%Y-%m-%d')

    # Save overview
    if client_data.get("overview"):
        with open(client_dir / "overview.md", 'w', encoding='utf-8') as f:
            f.write("---\n")
            f.write(f"title: \"{client_name} - Overview\"\n")
            f.write(f"type: \"Client Overview\"\n")
            f.write(f"client: \"{client_name}\"\n")
            f.write(f"date_updated: \"{date_str}\"\n")
            f.write("---\n\n")
            f.write(f"# {client_name} - Overview\n\n")
            f.write(client_data["overview"])
        print(f"      [OK] Saved: {client_dir / 'overview.md'}")

    # Save locations
    if client_data.get("locations"):
        with open(client_dir / "locations.md", 'w', encoding='utf-8') as f:
            f.write("---\n")
            f.write(f"title: \"{client_name} - Locations\"\n")
            f.write(f"type: \"Client Locations\"\n")
            f.write(f"client: \"{client_name}\"\n")
            f.write(f"date_updated: \"{date_str}\"\n")
            f.write("---\n\n")
            f.write(f"# {client_name} - Locations\n\n")
            for loc in client_data["locations"]:
                f.write(f"- {loc}\n")
        print(f"      [OK] Saved: {client_dir / 'locations.md'}")

    # Save hardware
    if client_data.get("hardware"):
        with open(client_dir / "hardware.md", 'w', encoding='utf-8') as f:
            f.write("---\n")
            f.write(f"title: \"{client_name} - Hardware\"\n")
            f.write(f"type: \"Client Hardware\"\n")
            f.write(f"client: \"{client_name}\"\n")
            f.write(f"date_updated: \"{date_str}\"\n")
            f.write("---\n\n")
            f.write(f"# {client_name} - Hardware\n\n")
            for hw in client_data["hardware"]:
                f.write(f"- {hw}\n")
        print(f"      [OK] Saved: {client_dir / 'hardware.md'}")

    # Save configuration
    if client_data.get("configuration"):
        with open(client_dir / "configuration.md", 'w', encoding='utf-8') as f:
            f.write("---\n")
            f.write(f"title: \"{client_name} - Configuration\"\n")
            f.write(f"type: \"Client Configuration\"\n")
            f.write(f"client: \"{client_name}\"\n")
            f.write(f"date_updated: \"{date_str}\"\n")
            f.write("---\n\n")
            f.write(f"# {client_name} - Configuration\n\n")
            for cfg in client_data["configuration"]:
                f.write(f"- {cfg}\n")
        print(f"      [OK] Saved: {client_dir / 'configuration.md'}")

    # Save contacts
    if client_data.get("contacts"):
        with open(client_dir / "contacts.md", 'w', encoding='utf-8') as f:
            f.write("---\n")
            f.write(f"title: \"{client_name} - Contacts\"\n")
            f.write(f"type: \"Client Contacts\"\n")
            f.write(f"client: \"{client_name}\"\n")
            f.write(f"date_updated: \"{date_str}\"\n")
            f.write("---\n\n")
            f.write(f"# {client_name} - Contacts\n\n")
            for contact in client_data["contacts"]:
                f.write(f"- {contact}\n")
        print(f"      [OK] Saved: {client_dir / 'contacts.md'}")

    # Save any additional dynamic categories
    for key, value in client_data.items():
        if key not in ['overview', 'locations', 'hardware', 'configuration', 'contacts']:
            filename = f"{key}.md"
            with open(client_dir / filename, 'w', encoding='utf-8') as f:
                f.write("---\n")
                f.write(f"title: \"{client_name} - {key.title()}\"\n")
                f.write(f"type: \"Client {key.title()}\"\n")
                f.write(f"client: \"{client_name}\"\n")
                f.write(f"date_updated: \"{date_str}\"\n")
                f.write("---\n\n")
                f.write(f"# {client_name} - {key.title()}\n\n")
                if isinstance(value, list):
                    for item in value:
                        f.write(f"- {item}\n")
                else:
                    f.write(str(value))
            print(f"      [OK] Saved: {client_dir / filename}")


def find_docx_files(directory: str, recursive: bool = False) -> List[Path]:
    """Find all .docx files in a directory

    Args:
        directory: Directory path to scan
        recursive: If True, scan subdirectories recursively. If False, only scan top-level.
    """
    directory_path = Path(directory)
    if not directory_path.exists():
        return []

    # Find all .docx files, excluding temporary files (starting with ~$)
    if recursive:
        # Recursive scan using rglob
        docx_files = [
            f for f in directory_path.rglob("*.docx")
            if not f.name.startswith("~$")
        ]
    else:
        # Non-recursive scan using glob
        docx_files = [
            f for f in directory_path.glob("*.docx")
            if not f.name.startswith("~$")
        ]

    return sorted(docx_files)


def process_document(doc_path: str, base_dir: Path, input_dir: str, client: OpenAI) -> bool:
    """Process a single document and add to knowledge base. Returns True if successful."""
    try:
        print("\n" + "=" * 70)
        print(f"[Document] Processing: {Path(doc_path).name}")
        print("=" * 70)

        # Scan existing structure from input directory
        existing_structure = scan_existing_structure(input_dir)

        # Extract document structure
        try:
            doc_data = extract_document_structure(doc_path)
            content = doc_data["full_content"]  # For backward compatibility with metadata extraction
            doc_metadata = doc_data.get("metadata", {})  # Get document metadata
        except Exception as e:
            print(f"[ERROR] Error reading document: {e}")
            return False

        # Extract metadata (products, client, type, category)
        metadata = extract_metadata(content, client, existing_structure)
        products = metadata.get("products", [])
        client_name = metadata.get("client_name", "Unknown")
        doc_type = metadata.get("document_type", "Specification")
        doc_category = metadata.get("document_category", "General")

        print(f"[Analysis] Document Analysis:")
        print(f"   Products: {', '.join(products)}")
        print(f"   Client: {client_name}")
        print(f"   Type: {doc_type}")
        print(f"   Category: {doc_category}\n")

        # Process each product
        for product in products:
            print(f"[Processing] Product: {product}")

            # Extract product knowledge
            knowledge = extract_product_knowledge(product, doc_data, client, existing_structure)

            # Check if there's sufficient information
            if knowledge and "INSUFFICIENT_INFORMATION" in knowledge:
                print(f"   [SKIP] {product} - insufficient information in document\n")
                continue

            # Extract reference materials (guides, templates, procedures, etc.)
            reference_materials = extract_document_template(product, doc_data, client, existing_structure)

            # Save everything with document metadata
            print(f"   [Saving] Files for {product}...")
            save_product_knowledge(base_dir, product, knowledge, reference_materials, doc_type, doc_category, doc_metadata)

        # Extract and save client information
        if client_name and client_name.lower() not in ['none', 'unknown', 'n/a']:
            print(f"[Processing] Client: {client_name}")
            client_data = extract_client_info(client_name, content, client, existing_structure)
            if client_data:
                print(f"   [Saving] Files for {client_name}...")
                save_client_info(base_dir, client_name, client_data)

        print(f"\n[SUCCESS] Completed: {Path(doc_path).name}")
        return True

    except Exception as e:
        print(f"\n[ERROR] Failed to process document: {e}")
        import traceback
        traceback.print_exc()
        return False


# ============================================================================
# Main Function
# ============================================================================

def main():
    """Main entry point"""
    parser = argparse.ArgumentParser(
        description="Knowledge Base Builder - Extract knowledge from Word documents using Kimi 2.5",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Process a single document
  python knowledge_base_builder_kimi.py document.docx -o knowledge_base

  # Process all documents in a directory
  python knowledge_base_builder_kimi.py -d ./documents -o knowledge_base

  # Process all documents recursively
  python knowledge_base_builder_kimi.py -d ./documents -r -o knowledge_base

  # Process a specific document by number from directory listing
  python knowledge_base_builder_kimi.py -d ./documents -o knowledge_base
  (then enter the document number when prompted)
        """
    )

    parser.add_argument(
        'document',
        nargs='?',
        help='Path to Word document (.docx) to process'
    )
    parser.add_argument(
        '-d', '--dir',
        help='Process all .docx files in this directory'
    )
    parser.add_argument(
        '-r', '--recursive',
        action='store_true',
        help='Scan directory recursively for .docx files'
    )
    parser.add_argument(
        '-o', '--output',
        default='knowledge_base',
        help='Output directory for knowledge base (default: knowledge_base)'
    )

    args = parser.parse_args()

    # Validate arguments
    if not args.document and not args.dir:
        parser.error("Either provide a document path or use --dir to process a directory")

    # Debug: Show environment variables
    print("\n[DEBUG] Environment Configuration:")
    print(f"  MOONSHOT_API_KEY: {'SET (' + MOONSHOT_API_KEY[:8] + '...' + MOONSHOT_API_KEY[-4:] + ')' if MOONSHOT_API_KEY else 'NOT SET'}")
    print(f"  MOONSHOT_BASE_URL: {MOONSHOT_BASE_URL}")
    print(f"  API Key Length: {len(MOONSHOT_API_KEY) if MOONSHOT_API_KEY else 0} characters")
    print()

    # Validate API key
    if not MOONSHOT_API_KEY:
        print("[ERROR] MOONSHOT_API_KEY environment variable is not set")
        print("Please set it with: export MOONSHOT_API_KEY=your_api_key")
        print("Or on Windows: set MOONSHOT_API_KEY=your_api_key")
        return 1

    # Initialize OpenAI client for Moonshot API
    client = OpenAI(
        api_key=MOONSHOT_API_KEY,
        base_url=MOONSHOT_BASE_URL
    )

    # Test API connection
    try:
        print("[Initializing] Testing Moonshot API connection...")
        client.chat.completions.create(
            model=MODEL_FAST,
            messages=[{"role": "user", "content": "Hello"}],
            max_tokens=10
        )
        print("[OK] Moonshot API connection successful\n")
    except Exception as e:
        print(f"[ERROR] Failed to connect to Moonshot API: {e}")
        print("Please check your API key and internet connection.")
        return 1

    base_dir = Path(args.output)
    base_dir.mkdir(parents=True, exist_ok=True)

    # Process documents
    if args.dir:
        # Directory mode
        docx_files = find_docx_files(args.dir, args.recursive)

        if not docx_files:
            print(f"[ERROR] No .docx files found in {args.dir}")
            return 1

        print(f"[Scanning] Found {len(docx_files)} document(s) to process:")
        for idx, doc_file in enumerate(docx_files, 1):
            print(f"   {idx}. {doc_file.name}")

        # Ask user for confirmation with option to select specific document
        response = input(f"\nProcess all {len(docx_files)} document(s)? (y=yes all, n=no, e=each, or enter number): ").strip().lower()

        if response == 'n':
            print("[Cancelled] No documents processed")
            return 0
        elif response.isdigit():
            # User entered a number - process only that document
            doc_num = int(response)
            if 1 <= doc_num <= len(docx_files):
                selected_doc_index = doc_num - 1
                print(f"[OK] Processing only document #{doc_num}: {docx_files[selected_doc_index].name}\n")
                success = process_document(str(docx_files[selected_doc_index]), base_dir, args.output, client)
                if success:
                    print(f"\n[COMPLETE] Successfully processed 1 document")
                    print(f"[KB] Knowledge base location: {base_dir.absolute()}")
                    return 0
                else:
                    print(f"\n[FAILED] Document processing failed")
                    return 1
            else:
                print(f"[ERROR] Invalid document number. Please enter a number between 1 and {len(docx_files)}")
                return 1
        elif response == 'e':
            # Process each with confirmation
            processed = 0
            for doc_file in docx_files:
                confirm = input(f"\nProcess {doc_file.name}? (y/n): ").strip().lower()
                if confirm == 'y':
                    success = process_document(str(doc_file), base_dir, args.output, client)
                    if success:
                        processed += 1
            print(f"\n[COMPLETE] Successfully processed {processed}/{len(docx_files)} documents")
            print(f"[KB] Knowledge base location: {base_dir.absolute()}")
        else:
            # Process all
            processed = 0
            for doc_file in docx_files:
                success = process_document(str(doc_file), base_dir, args.output, client)
                if success:
                    processed += 1

            print(f"\n[COMPLETE] Successfully processed {processed}/{len(docx_files)} documents")
            print(f"[KB] Knowledge base location: {base_dir.absolute()}")

    else:
        # Single document mode
        doc_path = args.document
        if not Path(doc_path).exists():
            print(f"[ERROR] Document not found: {doc_path}")
            return 1

        success = process_document(doc_path, base_dir, args.output, client)
        if success:
            print(f"\n[COMPLETE] Successfully processed document")
            print(f"[KB] Knowledge base location: {base_dir.absolute()}")
            return 0
        else:
            return 1

    return 0


if __name__ == "__main__":
    exit(main())

