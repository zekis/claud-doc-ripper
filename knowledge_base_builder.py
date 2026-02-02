#!/usr/bin/env python3
"""
Knowledge Base Builder - Builds cumulative knowledge base from Word documents

Organizes extracted information into:
- Products/[ProductName]/ - Product knowledge that accumulates over time
- Clients/[ClientName]/ - Client information split into categories
- Reference Materials - Guides, templates, procedures, and other reusable content organized by product and type
"""

import sys
import os
import json
import argparse
from pathlib import Path
from typing import List
from docx import Document
from auggie_sdk import Auggie


def find_auggie_cli():
    """Find auggie CLI on Windows"""
    import shutil

    # Try standard which
    cli_path = shutil.which("auggie")
    if cli_path:
        return cli_path

    # Try Windows npm paths
    if sys.platform == "win32":
        possible_paths = [
            os.path.expandvars(r"%APPDATA%\npm\auggie.cmd"),
            os.path.expandvars(r"%APPDATA%\npm\auggie.CMD"),
        ]
        for path in possible_paths:
            if os.path.exists(path):
                return path

    return None


def scan_existing_structure(output_dir: str) -> str:
    """Scan existing knowledge base structure and return formatted description"""
    output_path = Path(output_dir)

    if not output_path.exists():
        return "No existing knowledge base found. This is a new knowledge base."

    structure = []
    structure.append("EXISTING KNOWLEDGE BASE STRUCTURE:")
    structure.append("")

    # Scan Products
    products_dir = output_path / "Products"
    if products_dir.exists():
        structure.append("Products/")
        for product_dir in sorted(products_dir.iterdir()):
            if product_dir.is_dir():
                structure.append(f"  {product_dir.name}/")

                # Check for overview
                if (product_dir / "overview.md").exists():
                    structure.append(f"    overview.md (existing product knowledge)")

                # Check for reference materials
                reference_dir = product_dir / "Reference Materials"
                if reference_dir.exists():
                    structure.append(f"    Reference Materials/")
                    for category_dir in sorted(reference_dir.iterdir()):
                        if category_dir.is_dir():
                            structure.append(f"      {category_dir.name}/")
                            for ref_file in sorted(category_dir.glob("*.md")):
                                structure.append(f"        {ref_file.name}")

    # Scan Clients
    clients_dir = output_path / "Clients"
    if clients_dir.exists():
        structure.append("")
        structure.append("Clients/")
        for client_dir in sorted(clients_dir.iterdir()):
            if client_dir.is_dir():
                structure.append(f"  {client_dir.name}/")
                for client_file in sorted(client_dir.glob("*.md")):
                    structure.append(f"    {client_file.stem}.md")

    if len(structure) == 2:  # Only header
        return "No existing knowledge base found. This is a new knowledge base."

    return "\n".join(structure)


def read_word_document(doc_path: str) -> str:
    """Read Word document and extract text content"""
    print(f"📖 Reading: {doc_path}")
    doc = Document(doc_path)
    content = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name
            content.append(f"[{style}] {para.text}")
    
    # Extract tables
    for table in doc.tables:
        content.append("\n[TABLE]")
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            content.append(row_text)
        content.append("[/TABLE]\n")
    
    full_content = "\n".join(content)
    print(f"✅ Read {len(full_content)} characters\n")
    return full_content


def extract_metadata(content: str, cli_path: str, existing_structure: str) -> dict:
    """Extract metadata: products mentioned, client name, document type"""
    print("🔍 Extracting document metadata...")

    sample = content[:15000]
    agent = Auggie(model="haiku4.5", cli_path=cli_path)

    try:
        result = agent.run(
            f"""Analyze this document and extract metadata in JSON format.

            {existing_structure}

            DOCUMENT CONTENT:
            {sample}

            Extract:
            1. products: List of product names mentioned (e.g., ["Insight CM", "QMS", "SMS"])
               - If a product already exists in the structure above, use the EXACT same name
               - Only create new product names if they are genuinely new products
            2. client_name: The client/company name (e.g., "Roy Hill")
               - If a client already exists in the structure above, use the EXACT same name
            3. document_type: Type of document (e.g., "Functional Specification", "User Manual", "Technical Guide")
            4. document_category: Category (e.g., "Controls", "Reports", "Configuration")
               - If similar categories exist in the structure above, use consistent naming

            Return ONLY valid JSON in this exact format:
            {{
                "products": ["Product1", "Product2"],
                "client_name": "ClientName",
                "document_type": "DocumentType",
                "document_category": "Category"
            }}
            """,
            return_type=str,
            timeout=120
        )
        
        # Parse JSON
        metadata = json.loads(result)
        print(f"✅ Found: {len(metadata.get('products', []))} products, Client: {metadata.get('client_name', 'Unknown')}\n")
        return metadata
        
    except Exception as e:
        print(f"⚠️  Error extracting metadata: {e}")
        print(f"   Using defaults\n")
        return {
            "products": ["Unknown"],
            "client_name": "Unknown",
            "document_type": "Specification",
            "document_category": "General"
        }


def extract_product_knowledge(product_name: str, content: str, cli_path: str, existing_structure: str) -> str:
    """Extract knowledge about a specific product"""
    print(f"   📦 Extracting knowledge for: {product_name}")

    sample = content[:25000]
    agent = Auggie(model="sonnet4.5", cli_path=cli_path)

    try:
        result = agent.run(
            f"""Extract detailed information about {product_name} from this document.

            {existing_structure}

            NOTE: If {product_name} already exists in the structure above, this content will be ADDED to existing knowledge.
            Focus on extracting NEW information that complements what might already exist.

            DOCUMENT CONTENT:
            {sample}

            IMPORTANT: Return the actual markdown content directly. Do NOT say "I have created..." or describe what you're doing.
            Just return the knowledge article content itself.

            Create a comprehensive knowledge article about {product_name} with these sections:

            ## Overview
            - What is {product_name}?
            - Purpose and main capabilities

            ## Features
            - List all features and capabilities mentioned
            - Technical specifications

            ## Integration & Interfaces
            - How it integrates with other systems
            - Data exchanges and interfaces

            ## Configuration
            - Configuration options mentioned
            - Setup requirements

            ## Use Cases
            - How it's used in this context
            - Operational procedures

            Format as markdown. Be detailed and technical.
            Remove client-specific examples but keep general product information.
            Return ONLY the markdown content, nothing else.
            """,
            return_type=str,
            timeout=180
        )

        return result

    except Exception as e:
        print(f"      ❌ Error: {e}")
        return f"# {product_name}\n\nError extracting knowledge: {e}"


def extract_document_template(product_name: str, content: str, cli_path: str, existing_structure: str) -> list:
    """Extract reusable reference materials from document (can return multiple)"""
    print(f"   📋 Analyzing document for reusable reference materials: {product_name}")

    sample = content[:30000]
    agent = Auggie(model="sonnet4.5", cli_path=cli_path)

    try:
        result = agent.run(
            f"""We are an engineering company building a knowledge base. Analyze this document and determine ALL the ways we might use it in the future.

            {existing_structure}

            DOCUMENT CONTENT:
            {sample}

            TASK: Identify 1-3 different ways this document could be valuable as reference material.

            Consider these types of reference materials:
            - GUIDE (how to do something - instructional content)
            - TEMPLATE (structure for creating similar documents)
            - SPECIFICATION (technical details of a system)
            - PROCEDURE (step-by-step process)
            - REFERENCE (lookup information, standards, patterns)
            - BEST_PRACTICES (guidelines, recommendations, methodologies)

            For EACH valuable use case you identify, create a separate reference document.

            Return your answer as a JSON array with this structure:
            [
              {{
                "type": "GUIDE",
                "title": "How to Write Technical Specifications",
                "category": "Best Practices",
                "content": "# How to Write Technical Specifications\\n\\n[Full markdown content here...]"
              }},
              {{
                "type": "TEMPLATE",
                "title": "Specification Document Structure",
                "category": "Templates",
                "content": "# Specification Document Structure\\n\\n[Full markdown content here...]"
              }}
            ]

            Guidelines for extraction:
            - For GUIDES: Preserve instructional content, examples, methodology
            - For TEMPLATES: Extract structure, required/optional sections, formatting
            - For SPECIFICATIONS: Keep technical patterns, architecture, integration approaches
            - For PROCEDURES: Preserve process steps, decision points, roles
            - For REFERENCE: Organize lookup information, standards, patterns
            - For BEST_PRACTICES: Capture guidelines, recommendations, lessons learned

            IMPORTANT:
            - Remove client-specific details but keep the valuable patterns
            - Make each output actionable and reusable for {product_name}
            - Each "content" field should be complete markdown (not a summary)
            - Return ONLY valid JSON, nothing else
            - If only one use case, return array with one item
            """,
            return_type=str,
            timeout=180
        )

        # Parse JSON response
        import json
        reference_materials = json.loads(result)

        if not isinstance(reference_materials, list):
            reference_materials = [reference_materials]

        print(f"      ✅ Identified {len(reference_materials)} reference material(s)")
        return reference_materials

    except json.JSONDecodeError as e:
        print(f"      ❌ JSON Parse Error: {e}")
        print(f"      Raw response: {result[:200]}...")
        return [{
            "type": "REFERENCE",
            "title": f"{product_name} Reference",
            "category": "General",
            "content": f"# {product_name} Reference Material\n\nError parsing AI response: {e}\n\nRaw content:\n{result}"
        }]
    except Exception as e:
        print(f"      ❌ Error: {e}")
        return [{
            "type": "REFERENCE",
            "title": f"{product_name} Reference",
            "category": "General",
            "content": f"# {product_name} Reference Material\n\nError extracting content: {e}"
        }]


def extract_client_info(client_name: str, content: str, cli_path: str, existing_structure: str) -> dict:
    """Extract client information split into categories"""
    # Skip if no valid client name
    if not client_name or client_name.lower() in ['none', 'unknown', 'n/a']:
        return {}

    print(f"   👤 Extracting information for client: {client_name}")

    sample = content[:20000]
    agent = Auggie(model="haiku4.5", cli_path=cli_path)

    try:
        result = agent.run(
            f"""Extract information about {client_name} from this document and organize into categories.

            {existing_structure}

            NOTE: Check if {client_name} already exists in the structure above.
            If they do, check what categories already exist (e.g., overview.md, locations.md, hardware.md).
            You can use existing categories OR suggest new ones that fit the pattern.

            DOCUMENT CONTENT:
            {sample}

            Return ONLY valid JSON in this exact format:
            {{
                "overview": "Brief overview of the client and project",
                "locations": ["Location 1 with details", "Location 2 with details"],
                "hardware": ["Hardware item 1", "Hardware item 2"],
                "configuration": ["Config detail 1", "Config detail 2"],
                "contacts": ["Contact 1", "Contact 2"]
            }}

            You can add additional categories if needed (e.g., "software", "network", "security").
            Be comprehensive and extract all relevant details.
            """,
            return_type=str,
            timeout=120
        )

        # Parse JSON
        client_data = json.loads(result)
        return client_data

    except Exception as e:
        print(f"      ❌ Error: {e}")
        return {
            "overview": f"Error extracting overview: {e}",
            "locations": [],
            "hardware": [],
            "configuration": [],
            "contacts": []
        }


def save_product_knowledge(base_dir: Path, product_name: str, knowledge: str,
                          reference_materials: list, doc_type: str, doc_category: str):
    """Save product knowledge and reference materials to appropriate folders"""
    product_dir = base_dir / "Products" / product_name
    product_dir.mkdir(parents=True, exist_ok=True)

    # Save product knowledge
    knowledge_file = product_dir / "overview.md"
    with open(knowledge_file, 'w', encoding='utf-8') as f:
        f.write(f"# {product_name}\n\n")
        f.write(f"*Last updated from document processing*\n\n")
        f.write(knowledge)
    print(f"      ✅ Saved: {knowledge_file}")

    # Save each reference material (guides, templates, procedures, etc.)
    for ref_material in reference_materials:
        ref_type = ref_material.get('type', 'REFERENCE')
        ref_title = ref_material.get('title', doc_type)
        ref_category = ref_material.get('category', doc_category)
        ref_content = ref_material.get('content', '')

        # Create safe filename from title
        safe_filename = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in ref_title)
        safe_filename = safe_filename.replace(' ', '_')

        reference_dir = product_dir / "Reference Materials" / ref_category
        reference_dir.mkdir(parents=True, exist_ok=True)
        reference_file = reference_dir / f"{safe_filename}.md"

        with open(reference_file, 'w', encoding='utf-8') as f:
            f.write(f"*Type: {ref_type} | Source: {doc_type}*\n\n")
            f.write(ref_content)
        print(f"      ✅ Saved: {reference_file}")


def save_client_info(base_dir: Path, client_name: str, client_data: dict):
    """Save client information split into separate files"""
    client_dir = base_dir / "Clients" / client_name
    client_dir.mkdir(parents=True, exist_ok=True)

    # Save overview
    if client_data.get("overview"):
        with open(client_dir / "overview.md", 'w', encoding='utf-8') as f:
            f.write(f"# {client_name} - Overview\n\n")
            f.write(client_data["overview"])
        print(f"      ✅ Saved: {client_dir / 'overview.md'}")

    # Save locations
    if client_data.get("locations"):
        with open(client_dir / "locations.md", 'w', encoding='utf-8') as f:
            f.write(f"# {client_name} - Locations\n\n")
            for loc in client_data["locations"]:
                f.write(f"- {loc}\n")
        print(f"      ✅ Saved: {client_dir / 'locations.md'}")

    # Save hardware
    if client_data.get("hardware"):
        with open(client_dir / "hardware.md", 'w', encoding='utf-8') as f:
            f.write(f"# {client_name} - Hardware\n\n")
            for hw in client_data["hardware"]:
                f.write(f"- {hw}\n")
        print(f"      ✅ Saved: {client_dir / 'hardware.md'}")

    # Save configuration
    if client_data.get("configuration"):
        with open(client_dir / "configuration.md", 'w', encoding='utf-8') as f:
            f.write(f"# {client_name} - Configuration\n\n")
            for cfg in client_data["configuration"]:
                f.write(f"- {cfg}\n")
        print(f"      ✅ Saved: {client_dir / 'configuration.md'}")

    # Save contacts
    if client_data.get("contacts"):
        with open(client_dir / "contacts.md", 'w', encoding='utf-8') as f:
            f.write(f"# {client_name} - Contacts\n\n")
            for contact in client_data["contacts"]:
                f.write(f"- {contact}\n")
        print(f"      ✅ Saved: {client_dir / 'contacts.md'}")


def find_docx_files(directory: str, recursive: bool = False) -> List[Path]:
    """Find all .docx files in a directory

    Args:
        directory: Directory path to scan
        recursive: If True, scan subdirectories recursively. If False, only scan top-level.
    """
    directory_path = Path(directory)
    if not directory_path.exists():
        return []

    # Find all .docx files, excluding temporary files (starting with ~$)
    if recursive:
        # Recursive scan using rglob
        docx_files = [
            f for f in directory_path.rglob("*.docx")
            if not f.name.startswith("~$")
        ]
    else:
        # Non-recursive scan using glob
        docx_files = [
            f for f in directory_path.glob("*.docx")
            if not f.name.startswith("~$")
        ]

    return sorted(docx_files)


def process_document(doc_path: str, base_dir: Path, input_dir: str, cli_path: str) -> bool:
    """Process a single document and add to knowledge base. Returns True if successful."""
    try:
        print("\n" + "=" * 70)
        print(f"📄 Processing: {Path(doc_path).name}")
        print("=" * 70)

        # Scan existing structure from input directory
        existing_structure = scan_existing_structure(input_dir)

        # Read document
        try:
            content = read_word_document(doc_path)
        except Exception as e:
            print(f"❌ Error reading document: {e}")
            return False

        # Extract metadata
        metadata = extract_metadata(content, cli_path, existing_structure)
        products = metadata.get("products", [])
        client_name = metadata.get("client_name", "Unknown")
        doc_type = metadata.get("document_type", "Specification")
        doc_category = metadata.get("document_category", "General")

        print(f"📊 Document Analysis:")
        print(f"   Products: {', '.join(products)}")
        print(f"   Client: {client_name}")
        print(f"   Type: {doc_type}")
        print(f"   Category: {doc_category}\n")

        # Process each product
        print(f"🔧 Processing {len(products)} product(s)...\n")
        for product in products:
            print(f"📦 Product: {product}")

            # Extract product knowledge
            knowledge = extract_product_knowledge(product, content, cli_path, existing_structure)

            # Extract reusable reference materials (can return multiple)
            reference_materials = extract_document_template(product, content, cli_path, existing_structure)

            # Save to knowledge base
            save_product_knowledge(base_dir, product, knowledge, reference_materials, doc_type, doc_category)
            print()

        # Process client information
        if client_name and client_name.lower() not in ['none', 'unknown', 'n/a']:
            print(f"👥 Processing client: {client_name}")
            client_data = extract_client_info(client_name, content, cli_path, existing_structure)
            save_client_info(base_dir, client_name, client_data)
            print()
        else:
            print(f"⏭️  Skipping client extraction (no client identified)")
            print()

        # Summary
        print("✅ Document processed successfully!")
        print(f"   Products: {', '.join(products)}")
        print(f"   Client: {client_name}")

        return True

    except Exception as e:
        print(f"❌ Error processing document: {e}")
        return False


def main():
    # Parse command-line arguments
    parser = argparse.ArgumentParser(
        description='Extract knowledge from Word documents into a cumulative knowledge base'
    )
    parser.add_argument(
        'document',
        nargs='?',
        help='Path to the Word document (.docx) to process'
    )
    parser.add_argument(
        '-d', '--dir',
        help='Directory to scan for .docx files'
    )
    parser.add_argument(
        '-r', '--recursive',
        action='store_true',
        help='Recursively scan subdirectories when using --dir'
    )
    parser.add_argument(
        '-y', '--yes',
        action='store_true',
        help='Auto-approve all documents without prompting (batch mode only)'
    )
    parser.add_argument(
        '-o', '--output',
        default='wiki_output',
        help='Output directory for the knowledge base (default: wiki_output)'
    )
    parser.add_argument(
        '-i', '--input',
        default=None,
        help='Input directory to scan for existing structure (default: same as output)'
    )

    args = parser.parse_args()

    # Validate arguments
    if not args.document and not args.dir:
        parser.error("Either provide a document path or use --dir to scan a directory")

    if args.document and args.dir:
        parser.error("Cannot specify both a document and --dir. Choose one.")

    # If input not specified, use output directory
    input_dir = args.input if args.input else args.output

    # Find auggie CLI
    cli_path = find_auggie_cli()
    if not cli_path:
        print("❌ Could not find auggie CLI")
        print("   Install with: npm install -g @augmentcode/auggie@prerelease")
        return 1

    print(f"✅ Found auggie CLI: {cli_path}\n")

    # Output directory
    base_dir = Path(args.output)

    # Determine mode: single document or directory scan
    if args.dir:
        # Directory mode
        print("=" * 70)
        print("📚 KNOWLEDGE BASE BUILDER - BATCH MODE")
        print("=" * 70)
        print(f"Scan Directory: {Path(args.dir).absolute()}")
        print(f"Recursive: {'Yes' if args.recursive else 'No (top-level only)'}")
        print(f"Input:  {Path(input_dir).absolute()}")
        print(f"Output: {base_dir.absolute()}")
        print("=" * 70 + "\n")

        # Find all .docx files
        scan_mode = "recursively" if args.recursive else "in top-level directory"
        print(f"🔍 Scanning {scan_mode} for .docx files...")
        docx_files = find_docx_files(args.dir, recursive=args.recursive)

        if not docx_files:
            print(f"❌ No .docx files found in {args.dir}")
            return 1

        print(f"✅ Found {len(docx_files)} document(s) to process:\n")
        for i, doc in enumerate(docx_files, 1):
            print(f"   {i}. {doc.name}")
        print()

        # Approval process
        approve_all = args.yes  # If --yes flag is set, auto-approve all
        approve_each = False

        if not approve_all:
            # Ask user for approval
            print("=" * 70)
            response = input(f"Process all {len(docx_files)} document(s)? (y=yes all, n=no, e=each): ").strip().lower()
            print("=" * 70 + "\n")

            if response in ['n', 'no']:
                print("❌ Processing cancelled by user.")
                return 0
            elif response in ['e', 'each']:
                approve_each = True
                print("📋 You will be prompted for each document.\n")
            elif response in ['y', 'yes']:
                approve_all = True
                print("✅ Processing all documents...\n")
            else:
                print("❌ Invalid response. Processing cancelled.")
                return 0

        # Process each document
        successful = 0
        failed = 0
        skipped = 0

        for i, doc_path in enumerate(docx_files, 1):
            # If approve_each is True, ask for confirmation for this document
            if approve_each:
                print(f"\n{'=' * 70}")
                print(f"📄 Document {i}/{len(docx_files)}: {doc_path.name}")
                print(f"{'=' * 70}")
                confirm = input("Process this document? (y/n/q=quit): ").strip().lower()

                if confirm in ['q', 'quit']:
                    print("\n❌ Processing stopped by user.")
                    break
                elif confirm not in ['y', 'yes']:
                    print("⏭️  Skipped")
                    skipped += 1
                    continue

            print(f"\n{'=' * 70}")
            print(f"📄 Document {i}/{len(docx_files)}")
            print(f"{'=' * 70}")

            if process_document(str(doc_path), base_dir, input_dir, cli_path):
                successful += 1
            else:
                failed += 1

        # Final summary
        print("\n" + "=" * 70)
        print("🎉 BATCH PROCESSING COMPLETE!")
        print("=" * 70)
        print(f"✅ Successful: {successful}")
        if failed > 0:
            print(f"❌ Failed: {failed}")
        if skipped > 0:
            print(f"⏭️  Skipped: {skipped}")
        print(f"📊 Total: {successful + failed + skipped}/{len(docx_files)}")
        print(f"📁 Knowledge Base: {base_dir.absolute()}")
        print("=" * 70)

    else:
        # Single document mode
        doc_path = args.document

        print("=" * 70)
        print("📚 KNOWLEDGE BASE BUILDER")
        print("=" * 70)
        print(f"Document: {doc_path}")
        print(f"Input:  {Path(input_dir).absolute()}")
        print(f"Output: {base_dir.absolute()}")
        print("=" * 70 + "\n")

        # Process the single document
        if not process_document(doc_path, base_dir, input_dir, cli_path):
            return 1

        print("\n" + "=" * 70)
        print("✅ KNOWLEDGE BASE UPDATED!")
        print("=" * 70)
        print(f"📁 Knowledge Base: {base_dir.absolute()}")
        print("=" * 70)

    return 0


if __name__ == "__main__":
    sys.exit(main())

